#!/usr/bin/env python3
"""
Gemini LaTeX Summarizer
=======================
Enterprise-grade document-to-LaTeX pipeline using Google Gemini AI.

Architecture: Strategy Pattern + Chain of Responsibility
License: MIT
"""

from __future__ import annotations

import base64
import mimetypes
import re
import sys
from abc import ABC, abstractmethod
from enum import Enum
from pathlib import Path
from typing import Protocol, runtime_checkable

import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document as DocxDocument
from PIL import Image
from pydantic import Field, field_validator
from pydantic_settings import BaseSettings, SettingsConfigDict
from rich.console import Console
from rich.panel import Panel
from rich.progress import (
    BarColumn,
    Progress,
    SpinnerColumn,
    TaskProgressColumn,
    TextColumn,
    TimeElapsedColumn,
)
from rich.syntax import Syntax
from tenacity import (
    retry,
    retry_if_exception_type,
    stop_after_attempt,
    wait_exponential,
)

# =============================================================================
# CONFIGURATION
# =============================================================================


class Settings(BaseSettings):
    """Type-safe configuration management via Pydantic."""

    model_config = SettingsConfigDict(
        env_file=".env",
        env_file_encoding="utf-8",
        extra="ignore",
    )

    gemini_api_key: str = Field(..., alias="GEMINI_API_KEY")
    model_name: str = Field(default="gemini-2.0-flash", alias="GEMINI_MODEL")
    max_chunk_size: int = Field(default=100_000, ge=1000)
    temperature: float = Field(default=0.3, ge=0.0, le=2.0)
    output_dir: Path = Field(default=Path("./output"))

    @field_validator("output_dir", mode="before")
    @classmethod
    def ensure_path(cls, v: str | Path) -> Path:
        path = Path(v)
        path.mkdir(parents=True, exist_ok=True)
        return path


# =============================================================================
# SYSTEM PROMPT - LaTeX Generation Instructions
# =============================================================================

SYSTEM_PROMPT = r"""Sei un assistente accademico specializzato nella redazione di documenti tecnici LaTeX di livello professionale.

## COMPITO
Analizza il documento fornito e genera un riassunto esaustivo in codice LaTeX puro, pronto per la compilazione su Overleaf.

## REGOLE IMPERATIVE
1. **Completezza**: NON omettere dettagli tecnici rilevanti. Ogni concetto chiave deve essere presente.
2. **Struttura gerarchica**: Organizza logicamente in \chapter{}, \section{}, \subsection{}.
3. **LaTeX avanzato**: Usa tcolorbox per definizioni/teoremi, booktabs per tabelle, enumitem per liste.
4. **Zero errori**: Il codice DEVE compilare immediatamente senza warning.
5. **Output puro**: Restituisci SOLO codice LaTeX, nessun commento esterno o markdown.

## STRUTTURA DOCUMENTO OBBLIGATORIA

```latex
\documentclass[11pt,a4paper,twoside]{scrreprt}

% ===== PACKAGES =====
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage[italian]{babel}
\usepackage{geometry}
\geometry{margin=2.5cm}
\usepackage{microtype}
\usepackage{lmodern}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{enumitem}
\usepackage{hyperref}
\hypersetup{colorlinks=true,linkcolor=blue!70!black,urlcolor=blue!50!black}
\usepackage{cleveref}
\usepackage[most]{tcolorbox}

% ===== TCOLORBOX DEFINITIONS =====
\newtcolorbox{keypoint}{
    colback=blue!5!white,
    colframe=blue!75!black,
    fonttitle=\bfseries,
    title=Punto Chiave,
    sharp corners,
    boxrule=0.5pt
}

\newtcolorbox{definition}{
    colback=green!5!white,
    colframe=green!50!black,
    fonttitle=\bfseries,
    title=Definizione,
    sharp corners,
    boxrule=0.5pt
}

\newtcolorbox{warning}{
    colback=red!5!white,
    colframe=red!75!black,
    fonttitle=\bfseries,
    title=Attenzione,
    sharp corners,
    boxrule=0.5pt
}

% ===== DOCUMENT =====
\begin{document}

\title{[TITOLO APPROPRIATO]}
\author{Generato con Gemini AI}
\date{\today}
\maketitle

\tableofcontents
\newpage

% --- Executive Summary ---
\chapter*{Executive Summary}
\addcontentsline{toc}{chapter}{Executive Summary}
[Sintesi di 200-300 parole dei punti salienti]

% --- Contenuto principale ---
\chapter{[Primo capitolo logico]}
[Contenuto dettagliato con sezioni e sottosezioni]

% --- Key Takeaways ---
\chapter*{Key Takeaways}
\addcontentsline{toc}{chapter}{Key Takeaways}
\begin{itemize}[leftmargin=*,itemsep=0.5em]
    \item [Punto 1]
    \item [Punto 2]
    \item [...]
\end{itemize}

\end{document}
```

## NOTE STILISTICHE
- Usa \emph{} per enfasi, \textbf{} con parsimonia
- Preferisci ambienti itemize/enumerate con customizzazioni enumitem
- Inserisci \label{} e \cref{} per riferimenti incrociati
- Tabelle: sempre con booktabs (\toprule, \midrule, \bottomrule)
"""


# =============================================================================
# FILE TYPE ENUMERATION
# =============================================================================


class FileType(Enum):
    """Supported file types enumeration."""

    PDF = "pdf"
    DOCX = "docx"
    IMAGE = "image"
    UNKNOWN = "unknown"

    @classmethod
    def from_path(cls, path: Path) -> FileType:
        """Detect file type from path extension and MIME type."""
        suffix = path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(path))

        if suffix == ".pdf" or mime_type == "application/pdf":
            return cls.PDF
        elif suffix == ".docx" or mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            return cls.DOCX
        elif suffix in (".jpg", ".jpeg", ".png", ".webp", ".gif") or (
            mime_type and mime_type.startswith("image/")
        ):
            return cls.IMAGE
        return cls.UNKNOWN


# =============================================================================
# EXTRACTION STRATEGY PATTERN
# =============================================================================


@runtime_checkable
class ContentExtractor(Protocol):
    """Protocol defining the extraction interface."""

    def extract(self, file_path: Path) -> tuple[str, list[bytes]]:
        """Extract text and images from document.

        Returns:
            Tuple of (extracted_text, list_of_image_bytes)
        """
        ...


class PDFExtractor:
    """PDF content extractor using PyMuPDF (fitz)."""

    def extract(self, file_path: Path) -> tuple[str, list[bytes]]:
        """Extract text and embedded images from PDF."""
        text_parts: list[str] = []
        images: list[bytes] = []

        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                # Extract text with layout preservation
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"--- Pagina {page_num} ---\n{text}")

                # Extract images (limit to avoid memory issues)
                if len(images) < 10:
                    for img_index, img in enumerate(page.get_images(full=True)):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            images.append(base_image["image"])
                        except Exception:
                            continue

        return "\n\n".join(text_parts), images


class DocxExtractor:
    """Word document (.docx) content extractor."""

    def extract(self, file_path: Path) -> tuple[str, list[bytes]]:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading styles
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    text_parts.append(f"\n{'#' * int(level)} {para.text}\n")
                else:
                    text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            if table_text:
                text_parts.append("\n[TABELLA]\n" + "\n".join(table_text) + "\n")

        return "\n".join(text_parts), []


class ImageExtractor:
    """Image file extractor with preprocessing."""

    MAX_DIMENSION = 2048

    def extract(self, file_path: Path) -> tuple[str, list[bytes]]:
        """Process image and return as bytes for multimodal input."""
        with Image.open(file_path) as img:
            # Convert to RGB if necessary
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")

            # Resize if too large
            if max(img.size) > self.MAX_DIMENSION:
                ratio = self.MAX_DIMENSION / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)

            # Save to bytes
            from io import BytesIO

            buffer = BytesIO()
            img.save(buffer, format="JPEG", quality=85)
            return "", [buffer.getvalue()]


class ExtractorFactory:
    """Factory for creating appropriate content extractor."""

    _extractors: dict[FileType, ContentExtractor] = {
        FileType.PDF: PDFExtractor(),
        FileType.DOCX: DocxExtractor(),
        FileType.IMAGE: ImageExtractor(),
    }

    @classmethod
    def get_extractor(cls, file_type: FileType) -> ContentExtractor:
        """Get extractor instance for file type."""
        if file_type not in cls._extractors:
            raise ValueError(f"No extractor available for file type: {file_type}")
        return cls._extractors[file_type]


# =============================================================================
# LATEX OUTPUT FORMATTER
# =============================================================================


def format_latex_output(raw_response: str) -> str:
    """Clean Gemini response to extract pure LaTeX code.

    Handles:
    - Markdown code blocks (```latex ... ```)
    - Leading/trailing whitespace
    - Common artifacts
    """
    text = raw_response.strip()

    # Remove markdown code blocks
    patterns = [
        r"```latex\s*(.*?)\s*```",
        r"```tex\s*(.*?)\s*```",
        r"```\s*(.*?)\s*```",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            text = match.group(1).strip()
            break

    # Ensure document structure exists
    if r"\documentclass" not in text:
        # Response might be partial, wrap it
        text = (
            r"\documentclass[11pt,a4paper]{scrreprt}" + "\n"
            r"\usepackage[utf8]{inputenc}" + "\n"
            r"\begin{document}" + "\n"
            + text + "\n"
            r"\end{document}"
        )

    return text


# =============================================================================
# GEMINI API CLIENT
# =============================================================================


class GeminiClient:
    """Wrapper for Google Gemini API with retry logic."""

    def __init__(self, settings: Settings):
        self.settings = settings
        genai.configure(api_key=settings.gemini_api_key)
        self.model = genai.GenerativeModel(
            model_name=settings.model_name,
            system_instruction=SYSTEM_PROMPT,
            generation_config=genai.GenerationConfig(
                temperature=settings.temperature,
                max_output_tokens=8192,
            ),
        )

    @retry(
        retry=retry_if_exception_type((Exception,)),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=60),
        reraise=True,
    )
    def generate_summary(
        self, text_content: str, images: list[bytes] | None = None
    ) -> str:
        """Generate LaTeX summary from content with retry logic."""
        parts = []

        # Add text content
        if text_content:
            parts.append(f"## CONTENUTO DEL DOCUMENTO\n\n{text_content}")

        # Add images as multimodal input
        if images:
            for idx, img_bytes in enumerate(images[:5]):  # Limit images
                parts.append({
                    "mime_type": "image/jpeg",
                    "data": base64.b64encode(img_bytes).decode("utf-8"),
                })
                parts.append(f"[Immagine {idx + 1} del documento]")

        # Generate response
        response = self.model.generate_content(parts)

        if not response.text:
            raise ValueError("Empty response from Gemini API")

        return response.text


# =============================================================================
# MAIN PIPELINE
# =============================================================================


class LaTeXSummarizer:
    """Main orchestrator for the document-to-LaTeX pipeline."""

    def __init__(self, settings: Settings | None = None):
        self.console = Console()
        self.settings = settings or Settings()
        self.client = GeminiClient(self.settings)

    def process_file(self, input_path: Path) -> Path:
        """Process a single file and generate LaTeX output."""
        input_path = Path(input_path).resolve()

        if not input_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")

        # Detect file type
        file_type = FileType.from_path(input_path)
        if file_type == FileType.UNKNOWN:
            raise ValueError(f"Unsupported file type: {input_path.suffix}")

        self.console.print(
            Panel(
                f"[bold blue]Processing:[/] {input_path.name}\n"
                f"[bold]Type:[/] {file_type.value.upper()}",
                title="📄 Gemini LaTeX Summarizer",
                border_style="blue",
            )
        )

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            TimeElapsedColumn(),
            console=self.console,
        ) as progress:
            # Step 1: Extract content
            task = progress.add_task("[cyan]Extracting content...", total=3)

            extractor = ExtractorFactory.get_extractor(file_type)
            text_content, images = extractor.extract(input_path)
            progress.update(task, advance=1)

            # Step 2: Generate summary via Gemini
            progress.update(task, description="[yellow]Generating LaTeX via Gemini...")
            raw_latex = self.client.generate_summary(text_content, images)
            progress.update(task, advance=1)

            # Step 3: Format output
            progress.update(task, description="[green]Formatting output...")
            clean_latex = format_latex_output(raw_latex)
            progress.update(task, advance=1)

        # Save output
        output_filename = input_path.stem + "_summary.tex"
        output_path = self.settings.output_dir / output_filename

        output_path.write_text(clean_latex, encoding="utf-8")

        # Display result
        self.console.print(
            Panel(
                Syntax(clean_latex[:2000] + "\n\n[...]" if len(clean_latex) > 2000 else clean_latex,
                       "latex", theme="monokai", line_numbers=True),
                title="📝 Generated LaTeX (preview)",
                border_style="green",
            )
        )

        self.console.print(
            f"\n[bold green]✓[/] Output saved to: [underline]{output_path}[/]\n"
        )

        return output_path


# =============================================================================
# CLI ENTRY POINT
# =============================================================================


def main() -> int:
    """CLI entry point."""
    console = Console()

    if len(sys.argv) < 2:
        console.print(
            Panel(
                "[bold]Usage:[/] python gemini_latex_summarizer.py <file_path>\n\n"
                "[bold]Supported formats:[/]\n"
                "  • PDF (.pdf)\n"
                "  • Word (.docx)\n"
                "  • Images (.jpg, .png, .webp)\n\n"
                "[bold]Configuration:[/]\n"
                "  Create a .env file with:\n"
                "  GEMINI_API_KEY=your_api_key_here",
                title="📚 Gemini LaTeX Summarizer",
                border_style="yellow",
            )
        )
        return 1

    input_file = Path(sys.argv[1])

    try:
        summarizer = LaTeXSummarizer()
        summarizer.process_file(input_file)
        return 0

    except FileNotFoundError as e:
        console.print(f"[bold red]Error:[/] {e}")
        return 1

    except ValueError as e:
        console.print(f"[bold red]Validation Error:[/] {e}")
        return 1

    except Exception as e:
        console.print(f"[bold red]Unexpected Error:[/] {e}")
        console.print_exception()
        return 1


if __name__ == "__main__":
    sys.exit(main())
